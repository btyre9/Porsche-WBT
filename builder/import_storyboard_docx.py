"""
vo_segment_parser.py
====================
Extends the existing storyboard parser to split Voiceover fields into
named segments, generate WellSaid-ready file names, and produce a manifest.

Notation standard for IDs (in the storyboard Word doc):
----------------------------------------------------------------
Instead of a single "Voiceover:" field, authors write one field per segment:

    Voiceover-INTRO:       Welcome to this module on the Taycan battery system.
    Voiceover-CLICK-BatteryOverview:   The pack consists of 396 prismatic cells.
    Voiceover-TAB-ChargingModes:       Porsche offers three charging modes.
    Voiceover-STEP-02:     In step two, verify the cooling circuit pressure.

Supported trigger prefixes:
    INTRO       — plays on slide entry (always present)
    CLICK-<Label>   — plays when a hotspot / button is clicked
    TAB-<Label>     — plays when a tab or accordion item is revealed
    STEP-<N>        — plays when user advances to step N within a slide

Generated file name pattern:
    SLD-<CourseID>-<SlideNum>-<TriggerType>[-<Label>].mp3

Examples:
    SLD-CC01-001-INTRO.mp3
    SLD-CC01-001-CLICK-BatteryOverview.mp3
    SLD-CC01-001-TAB-ChargingModes.mp3
    SLD-CC01-001-STEP-02.mp3

Manifest (CSV) columns:
    FileName, SlideID, CourseID, TriggerType, Label, VoiceoverText
"""

from __future__ import annotations

import argparse
import csv
import re
from pathlib import Path
from typing import Any

# ---------------------------------------------------------------------------
# Re-use / extend the existing parser internals
# ---------------------------------------------------------------------------

KV_RE = re.compile(r"^([^:]+):\s*(.*)$")
SLIDE_HEADING_RE = re.compile(
    r"^(?:#+\s*)?(slide|screen|scene)\s*[-_ ]*\d+", re.IGNORECASE
)

# Matches Voiceover-INTRO, Voiceover-CLICK-Label, Voiceover-TAB-Label, Voiceover-STEP-02
VO_SEGMENT_RE = re.compile(
    r"^Voiceover-(INTRO|CLICK-(.+)|TAB-(.+)|STEP-(\d+))$",
    re.IGNORECASE,
)

# Trigger types we understand
VALID_TRIGGERS = {"INTRO", "CLICK", "TAB", "STEP"}


# ---------------------------------------------------------------------------
# Key canonicalisation — extends the original _canonical_key
# ---------------------------------------------------------------------------

def _canonical_key(raw_key: str) -> str:
    raw = raw_key.strip()
    compact = re.sub(r"[\s_-]+", "", raw).lower()

    direct = {
        "slideid": "Slide-ID",
        "templateid": "Template-ID",
        "template": "Template-ID",
        "slidetitle": "Slide-Title",
        "title": "Slide-Title",
        "audiovo": "Audio-VO",
        "audio": "Audio-VO",
        "voiceover": "Voiceover",
        "vo": "Voiceover",
        "script": "Voiceover",
        "narration": "Voiceover",
        "captiontext": "Caption-Text",
        "onscreentext": "On-Screen-Text",
        "animationintro": "Animation-Intro",
        "interactiontype": "Interaction-Type",
        "question": "Question",
        "correctanswer": "Correct-Answer",
        "quizgroup": "Quiz-Group",
        "subtitle": "Subtitle",
        "image": "Image",
    }
    if compact in direct:
        return direct[compact]

    # Voiceover segment fields — preserve exactly so the VO parser can find them
    # e.g. "voiceover-intro", "voiceover-click-batteryoverview"
    if compact.startswith("voiceover-"):
        # Reconstruct with normalised casing but keep the label
        parts = raw.split("-", 1)          # ["Voiceover", "CLICK-BatteryOverview"]
        if len(parts) == 2:
            trigger_part = parts[1].strip()
            sub = trigger_part.split("-", 1)
            trigger_type = sub[0].upper()
            if trigger_type in VALID_TRIGGERS:
                if len(sub) == 2:
                    label = _normalise_label(sub[1])
                    return f"Voiceover-{trigger_type}-{label}"
                return f"Voiceover-{trigger_type}"

    m_choice = re.match(r"^choice(\d+)$", compact)
    if m_choice:
        return f"Choice-{m_choice.group(1)}"

    m_objective = re.match(r"^objective(\d+)$", compact)
    if m_objective:
        return f"Objective-{m_objective.group(1)}"

    m_anim_element = re.match(r"^animationelement(.+)$", compact)
    if m_anim_element:
        suffix = m_anim_element.group(1).strip()
        return f"Animation-Element-{suffix}" if suffix else "Animation-Element-1"

    return raw.replace(" ", "-")


def _normalise_label(label: str) -> str:
    """Strip spaces, title-case — used for CLICK/TAB labels in file names."""
    return re.sub(r"\s+", "", label.strip().title())


# ---------------------------------------------------------------------------
# Document iteration (unchanged from original)
# ---------------------------------------------------------------------------

def _looks_like_slide_heading(text: str, style_name: str) -> bool:
    if SLIDE_HEADING_RE.match(text):
        return True
    if style_name.lower().startswith("heading") and text.lower().startswith("slide"):
        return True
    return False


def _append_value(slide: dict[str, Any], key: str, value: str) -> None:
    existing = slide.get(key, "")
    if existing:
        slide[key] = f"{existing} {value}".strip()
    else:
        slide[key] = value.strip()


def _iter_doc_lines(doc: Any) -> list[tuple[str, str]]:
    lines: list[tuple[str, str]] = []

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        style_name = para.style.name if para.style else ""
        for line in text.splitlines():
            clean = line.strip()
            if clean:
                lines.append((clean, style_name))

    for table in doc.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            cells = [c for c in cells if c]
            if not cells:
                continue
            if len(cells) >= 2:
                key_cell = cells[0]
                value_cell = " ".join(cells[1:]).strip()
                if key_cell.lower().startswith("slide") and not value_cell:
                    lines.append((key_cell, "Table"))
                    continue
                combined = (
                    f"{key_cell} {value_cell}".strip()
                    if ":" not in key_cell
                    else f"{key_cell}: {value_cell}".strip()
                )
                lines.append((combined, "Table"))
                continue
            only = cells[0]
            for line in only.splitlines():
                clean = line.strip()
                if clean:
                    lines.append((clean, "Table"))

    return lines


# ---------------------------------------------------------------------------
# Core extraction
# ---------------------------------------------------------------------------

def _extract_slides_from_docx(
    docx_path: Path,
) -> tuple[str, list[dict[str, Any]]]:
    try:
        from docx import Document
    except ImportError as exc:
        raise SystemExit(
            "python-docx is required.  Install: python -m pip install python-docx"
        ) from exc

    doc = Document(str(docx_path))
    course_title = ""
    slides: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None
    last_key: str | None = None

    def flush_current() -> None:
        nonlocal current, last_key
        if current and any(k for k in current.keys() if k != "section_heading"):
            slides.append(current)
        current = None
        last_key = None

    for line, style_name in _iter_doc_lines(doc):
        if line.lower().startswith("# course:"):
            course_title = line.split(":", 1)[1].strip()
            continue
        if line.lower().startswith("course:") and not course_title:
            course_title = line.split(":", 1)[1].strip()
            continue

        heading_text = line[2:].strip() if line.startswith("##") else line
        if _looks_like_slide_heading(heading_text, style_name):
            flush_current()
            current = {"section_heading": heading_text}
            continue

        m = KV_RE.match(line)
        if m:
            key = _canonical_key(m.group(1))
            value = m.group(2).strip()
            if key == "Slide-ID" and current and "Slide-ID" in current:
                flush_current()
                current = {"section_heading": f"Slide{len(slides) + 1:02d}"}

            if current is None:
                current = {"section_heading": f"Slide{len(slides) + 1:02d}"}
            current[key] = value
            last_key = key
            continue

        if current is None:
            continue
        if last_key:
            _append_value(current, last_key, line)

    flush_current()

    if not course_title:
        course_title = docx_path.stem.replace("-", " ").strip()

    for idx, slide in enumerate(slides, start=1):
        slide.setdefault("Slide-ID", f"slide-{idx:02d}")
        slide.setdefault("Template-ID", "content-standard-01")
        slide.setdefault("Slide-Title", slide.get("section_heading", f"Slide {idx}"))

    return course_title, slides


# ---------------------------------------------------------------------------
# VO segment extraction
# ---------------------------------------------------------------------------

def _parse_slide_id_parts(slide_id: str) -> tuple[str, str]:
    """
    Extract course_id and zero-padded slide_num from a Slide-ID like
    'CC01-003', 'cc01-003', or fall back gracefully.

    Returns (course_id, slide_num) e.g. ("CC01", "003")
    """
    slide_id = slide_id.strip().upper()
    # Try pattern like CC01-003
    m = re.match(r"^([A-Z]{2}\d+)-(\d+)$", slide_id)
    if m:
        return m.group(1), m.group(2).zfill(3)
    # Try just a number
    m2 = re.match(r"^(\d+)$", slide_id)
    if m2:
        return "COURSE", m2.group(1).zfill(3)
    # Fallback — treat whole thing as course_id, no slide_num
    return slide_id, "000"


def _build_file_name(course_id: str, slide_num: str, trigger_type: str, label: str) -> str:
    """
    Produce a file name following the convention:
        SLD-<CourseID>-<SlideNum>-<TriggerType>[-<Label>].mp3
    """
    parts = ["SLD", course_id, slide_num, trigger_type]
    if label:
        parts.append(label)
    return "-".join(parts) + ".mp3"


def _extract_vo_segments(
    slide: dict[str, Any],
    course_id: str,
    slide_num: str,
) -> list[dict[str, str]]:
    """
    Scan a parsed slide dict for all Voiceover-* fields and return a list of
    segment dicts ready for the manifest and WellSaid API.

    Also handles a plain 'Voiceover' field (no trigger tag) by treating it
    as INTRO so legacy storyboards don't break.
    """
    segments: list[dict[str, str]] = []

    for key, value in slide.items():
        if not value:
            continue

        # Legacy plain Voiceover field — treat as INTRO
        if key == "Voiceover":
            file_name = _build_file_name(course_id, slide_num, "INTRO", "")
            segments.append(
                {
                    "FileName": file_name,
                    "SlideID": slide.get("Slide-ID", ""),
                    "CourseID": course_id,
                    "TriggerType": "INTRO",
                    "Label": "",
                    "VoiceoverText": value.strip(),
                }
            )
            continue

        # Structured Voiceover-<TRIGGER>[-<Label>] fields
        if not key.startswith("Voiceover-"):
            continue

        remainder = key[len("Voiceover-"):]          # e.g. "INTRO", "CLICK-BatteryOverview"
        sub = remainder.split("-", 1)
        trigger_type = sub[0].upper()
        label = sub[1] if len(sub) == 2 else ""

        if trigger_type not in VALID_TRIGGERS:
            # Unknown trigger — skip with a warning
            print(f"  [WARNING] Unrecognised trigger '{trigger_type}' in key '{key}' — skipped.")
            continue

        # Zero-pad STEP labels
        if trigger_type == "STEP" and label.isdigit():
            label = label.zfill(2)

        file_name = _build_file_name(course_id, slide_num, trigger_type, label)
        segments.append(
            {
                "FileName": file_name,
                "SlideID": slide.get("Slide-ID", ""),
                "CourseID": course_id,
                "TriggerType": trigger_type,
                "Label": label,
                "VoiceoverText": value.strip(),
            }
        )

    # Preserve storyboard order: INTRO first, then CLICK/TAB/STEP
    trigger_order = {"INTRO": 0, "CLICK": 1, "TAB": 2, "STEP": 3}
    segments.sort(key=lambda s: trigger_order.get(s["TriggerType"], 99))
    return segments


# ---------------------------------------------------------------------------
# Manifest output
# ---------------------------------------------------------------------------

MANIFEST_FIELDS = ["FileName", "SlideID", "CourseID", "TriggerType", "Label", "VoiceoverText"]


def _write_manifest(
    segments: list[dict[str, str]], output_path: Path
) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with output_path.open("w", newline="", encoding="utf-8") as fh:
        writer = csv.DictWriter(fh, fieldnames=MANIFEST_FIELDS)
        writer.writeheader()
        writer.writerows(segments)


# ---------------------------------------------------------------------------
# Markdown output (mirrors original parser's _render_markdown)
# ---------------------------------------------------------------------------

def _order_slide_keys(slide: dict[str, Any]) -> list[str]:
    preferred = [
        "Slide-ID",
        "Template-ID",
        "Slide-Title",
        "Audio-VO",
        "Voiceover",
        "Caption-Text",
        "On-Screen-Text",
        "Animation-Intro",
        "Interaction-Type",
        "Question",
        "Correct-Answer",
        "Quiz-Group",
    ]
    keys = [k for k in slide.keys() if k != "section_heading"]
    keys_sorted: list[str] = []

    for key in preferred:
        if key in keys:
            keys_sorted.append(key)

    # Voiceover segment keys — sorted by trigger order, then label
    vo_keys = [k for k in keys if k.startswith("Voiceover-")]
    trigger_order = {"INTRO": 0, "CLICK": 1, "TAB": 2, "STEP": 3}
    vo_keys.sort(
        key=lambda k: (
            trigger_order.get(k.split("-")[1].upper() if "-" in k else "", 99),
            k,
        )
    )
    keys_sorted.extend(vo_keys)

    choice_keys = sorted(
        [k for k in keys if re.match(r"^Choice-\d+$", k)],
        key=lambda x: int(x.split("-")[1]),
    )
    keys_sorted.extend(choice_keys)

    objective_keys = sorted(
        [k for k in keys if re.match(r"^Objective-\d+$", k)],
        key=lambda x: int(x.split("-")[1]),
    )
    keys_sorted.extend(objective_keys)

    anim_element_keys = sorted(
        [k for k in keys if k.startswith("Animation-Element-")]
    )
    keys_sorted.extend(anim_element_keys)

    for key in keys:
        if key not in keys_sorted:
            keys_sorted.append(key)

    return keys_sorted


def _render_markdown(course_title: str, slides: list[dict[str, Any]]) -> str:
    out: list[str] = [f"# Course: {course_title}", ""]
    for idx, slide in enumerate(slides, start=1):
        section_heading = slide.get("section_heading", f"Slide{idx:02d}")
        out.append(f"## {section_heading}")
        for key in _order_slide_keys(slide):
            out.append(f"{key}: {slide.get(key, '')}")
        out.append("")
    return "\n".join(out).rstrip() + "\n"


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def run(
    docx_path: Path,
    output_markdown: Path,
    output_manifest: Path,
    course_id_override: str | None = None,
) -> None:
    if not docx_path.exists():
        raise FileNotFoundError(f"Word storyboard not found: {docx_path}")

    course_title, slides = _extract_slides_from_docx(docx_path)
    if not slides:
        raise ValueError(
            "No slides found in the Word document. "
            "Use slide headings like 'Slide01' and key/value lines like 'Slide-Title: ...'."
        )

    # Write markdown (unchanged behaviour from original parser)
    markdown = _render_markdown(course_title, slides)
    output_markdown.parent.mkdir(parents=True, exist_ok=True)
    output_markdown.write_text(markdown, encoding="utf-8")

    # Extract VO segments across all slides
    all_segments: list[dict[str, str]] = []
    for slide in slides:
        slide_id = slide.get("Slide-ID", "slide-00")

        if course_id_override:
            course_id = course_id_override.upper()
            # Try to pull slide number from Slide-ID
            m = re.search(r"(\d+)$", slide_id)
            slide_num = m.group(1).zfill(3) if m else "000"
        else:
            course_id, slide_num = _parse_slide_id_parts(slide_id)

        segments = _extract_vo_segments(slide, course_id, slide_num)
        all_segments.extend(segments)

    _write_manifest(all_segments, output_manifest)

    # Console summary
    print(f"\nStoryboard:  {docx_path}")
    print(f"Markdown:    {output_markdown}  ({len(slides)} slides)")
    print(f"Manifest:    {output_manifest}  ({len(all_segments)} VO segments)")
    print()

    if all_segments:
        print("VO segments found:")
        for seg in all_segments:
            trigger_label = f"{seg['TriggerType']}-{seg['Label']}" if seg["Label"] else seg["TriggerType"]
            preview = seg["VoiceoverText"][:60].replace("\n", " ")
            if len(seg["VoiceoverText"]) > 60:
                preview += "…"
            print(f"  {seg['FileName']:<55}  {trigger_label:<20}  \"{preview}\"")
    else:
        print(
            "No VO segments found. Make sure Voiceover fields use the notation:\n"
            "  Voiceover-INTRO:           <text>\n"
            "  Voiceover-CLICK-<Label>:   <text>\n"
            "  Voiceover-TAB-<Label>:     <text>\n"
            "  Voiceover-STEP-<N>:        <text>"
        )


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description=(
            "Parse a Word storyboard (.docx), write the standard markdown output, "
            "and produce a VO segment manifest for the WellSaid API."
        )
    )
    parser.add_argument(
        "--docx",
        type=Path,
        required=True,
        help="Path to source storyboard .docx file.",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("storyboard/course.md"),
        help="Output markdown path (default: storyboard/course.md).",
    )
    parser.add_argument(
        "--manifest",
        type=Path,
        default=Path("storyboard/vo_manifest.csv"),
        help="Output VO manifest CSV path (default: storyboard/vo_manifest.csv).",
    )
    parser.add_argument(
        "--course-id",
        type=str,
        default=None,
        help=(
            "Override the course ID used in file names (e.g. CC01). "
            "If omitted, it is inferred from each slide's Slide-ID field."
        ),
    )
    args = parser.parse_args()
    run(
        docx_path=args.docx.resolve(),
        output_markdown=args.output.resolve(),
        output_manifest=args.manifest.resolve(),
        course_id_override=args.course_id,
    )